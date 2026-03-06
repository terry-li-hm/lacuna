"""Gap analysis routes for Meridian API."""

import asyncio
import io
import logging
import os
from typing import Dict, Any
from fastapi import APIRouter, HTTPException, Depends
from fastapi.responses import Response
from jinja2 import Environment, FileSystemLoader
import weasyprint
from docx import Document
from docx.shared import RGBColor, Inches

from backend.state import (
    get_gap_analysis_service,
    get_document_service,
    get_policy_service,
)
from backend.models.schemas import (
    GapAnalysisRequest,
    GapAnalysisResponse,
    BatchGapAnalysisRequest,
    BatchGapAnalysisResponse,
    BatchGapAnalysisResult,
)

logger = logging.getLogger(__name__)
router = APIRouter()

# Setup Jinja2 environment
template_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "templates")
jinja_env = Environment(loader=FileSystemLoader(template_dir))

# In-memory cache: (circular_doc_id, baseline_id, is_policy_baseline, include_amendments) -> result
_gap_cache: Dict[tuple, Any] = {}


@router.post("/gap-analysis", response_model=GapAnalysisResponse)
async def gap_analysis(
    request: GapAnalysisRequest, service=Depends(get_gap_analysis_service)
):
    """
    Perform a gap analysis between a new circular and a baseline.
    Results are cached in-memory so repeat calls (e.g. during demos) return instantly.
    """
    logger.info(f"Gap Analysis: {request.circular_doc_id} vs {request.baseline_id}")

    cache_key = (
        request.circular_doc_id,
        request.baseline_id,
        request.is_policy_baseline,
        request.include_amendments,
    )
    if cache_key in _gap_cache and not request.no_llm:
        logger.info("Gap analysis cache hit — returning cached result")
        return _gap_cache[cache_key]

    try:
        result = await service.perform_gap_analysis(
            circular_doc_id=request.circular_doc_id,
            baseline_id=request.baseline_id,
            is_policy_baseline=request.is_policy_baseline,
            include_amendments=request.include_amendments,
            no_llm=request.no_llm,
        )
        if not request.no_llm:
            _gap_cache[cache_key] = result
        return result
    except ValueError as e:
        raise HTTPException(
            status_code=404 if "not found" in str(e).lower() else 400, detail=str(e)
        )
    except Exception as e:
        logger.error(f"Error performing gap analysis: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Internal server error")


@router.post("/gap-analysis/export")
async def export_gap_analysis(
    request: GapAnalysisRequest,
    service=Depends(get_gap_analysis_service),
    doc_service=Depends(get_document_service),
    policy_service=Depends(get_policy_service),
):
    """
    Perform a gap analysis and export the result as a Capco-branded PDF report.
    """
    logger.info(f"Exporting Gap Analysis: {request.circular_doc_id} vs {request.baseline_id}")

    cache_key = (
        request.circular_doc_id,
        request.baseline_id,
        request.is_policy_baseline,
        request.include_amendments,
    )
    
    if cache_key in _gap_cache and not request.no_llm:
        logger.info("Gap analysis cache hit — using cached result for export")
        result = _gap_cache[cache_key]
    else:
        try:
            result = await service.perform_gap_analysis(
                circular_doc_id=request.circular_doc_id,
                baseline_id=request.baseline_id,
                is_policy_baseline=request.is_policy_baseline,
                include_amendments=request.include_amendments,
                no_llm=request.no_llm,
            )
            if not request.no_llm:
                _gap_cache[cache_key] = result
        except ValueError as e:
            raise HTTPException(
                status_code=404 if "not found" in str(e).lower() else 400, detail=str(e)
            )
        except Exception as e:
            logger.error(f"Error performing gap analysis for export: {e}", exc_info=True)
            raise HTTPException(status_code=500, detail="Internal server error")

    # Get names for circular and baseline for the report
    circular_name = request.circular_doc_id
    baseline_name = request.baseline_id

    try:
        # Try to resolve circular name
        doc = doc_service.get_document(request.circular_doc_id)
        if doc:
            circular_name = doc.get("title") or doc.get("name") or request.circular_doc_id
            
        # Try to resolve baseline name
        if request.is_policy_baseline:
            policy = policy_service.get_policy(request.baseline_id)
            if policy:
                baseline_name = policy.get("name") or policy.get("title") or request.baseline_id
        else:
            doc = doc_service.get_document(request.baseline_id)
            if doc:
                baseline_name = doc.get("title") or doc.get("name") or request.baseline_id
    except Exception as e:
        logger.warning(f"Could not resolve circular or baseline names: {e}")

    # Prepare data for template
    template_data = {
        "circular_name": circular_name,
        "baseline_name": baseline_name,
        "generated_at": result.generated_at,
        "summary": result.summary,
        "findings": [f.model_dump() for f in result.findings],
    }

    try:
        # Render HTML
        template = jinja_env.get_template("gap_report.html")
        html_content = template.render(**template_data)

        # Generate PDF using WeasyPrint in a thread to avoid blocking event loop
        pdf_bytes = await asyncio.to_thread(
            lambda: weasyprint.HTML(string=html_content).write_pdf()
        )

        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="lacuna-gap-report.pdf"'
            },
        )
    except Exception as e:
        logger.error(f"Error generating PDF report: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to generate PDF report")


@router.post("/gap-analysis/export-docx")
async def export_gap_analysis_docx(
    request: GapAnalysisRequest,
    service=Depends(get_gap_analysis_service),
    doc_service=Depends(get_document_service),
    policy_service=Depends(get_policy_service),
):
    """
    Perform a gap analysis and export the result as a DOCX report.
    """
    logger.info(f"Exporting Gap Analysis to DOCX: {request.circular_doc_id} vs {request.baseline_id}")

    cache_key = (
        request.circular_doc_id,
        request.baseline_id,
        request.is_policy_baseline,
        request.include_amendments,
    )
    
    if cache_key in _gap_cache and not request.no_llm:
        logger.info("Gap analysis cache hit — using cached result for DOCX export")
        result = _gap_cache[cache_key]
    else:
        try:
            result = await service.perform_gap_analysis(
                circular_doc_id=request.circular_doc_id,
                baseline_id=request.baseline_id,
                is_policy_baseline=request.is_policy_baseline,
                include_amendments=request.include_amendments,
                no_llm=request.no_llm,
            )
            if not request.no_llm:
                _gap_cache[cache_key] = result
        except ValueError as e:
            raise HTTPException(
                status_code=404 if "not found" in str(e).lower() else 400, detail=str(e)
            )
        except Exception as e:
            logger.error(f"Error performing gap analysis for DOCX export: {e}", exc_info=True)
            raise HTTPException(status_code=500, detail="Internal server error")

    # Get names for circular and baseline for the report
    circular_name = request.circular_doc_id
    baseline_name = request.baseline_id

    try:
        # Try to resolve circular name
        doc = doc_service.get_document(request.circular_doc_id)
        if doc:
            circular_name = doc.get("title") or doc.get("name") or request.circular_doc_id
            
        # Try to resolve baseline name
        if request.is_policy_baseline:
            policy = policy_service.get_policy(request.baseline_id)
            if policy:
                baseline_name = policy.get("name") or policy.get("title") or request.baseline_id
        else:
            doc = doc_service.get_document(request.baseline_id)
            if doc:
                baseline_name = doc.get("title") or doc.get("name") or request.baseline_id
    except Exception as e:
        logger.warning(f"Could not resolve circular or baseline names for DOCX: {e}")

    def create_docx():
        document = Document()
        document.add_heading("Gap Analysis Report", level=0)

        # Metadata table
        table = document.add_table(rows=3, cols=2)
        table.cell(0, 0).text = "Circular"
        table.cell(0, 1).text = circular_name
        table.cell(1, 0).text = "Baseline"
        table.cell(1, 1).text = baseline_name
        table.cell(2, 0).text = "Generated Date"
        table.cell(2, 1).text = result.generated_at

        document.add_heading("Executive Summary", level=1)
        
        # Summary table
        summary_table = document.add_table(rows=1, cols=3)
        summary_table.style = 'Table Grid'
        
        hdr_cells = summary_table.rows[0].cells
        hdr_cells[0].text = f"Full Alignment: {result.summary.get('Full', 0)}"
        hdr_cells[1].text = f"Partial Alignment: {result.summary.get('Partial', 0)}"
        hdr_cells[2].text = f"Gap identified: {result.summary.get('Gap', 0)}"
        
        # Color fills
        for i, color in enumerate([RGBColor(0, 128, 0), RGBColor(218, 165, 32), RGBColor(255, 0, 0)]):
            if hdr_cells[i].paragraphs and hdr_cells[i].paragraphs[0].runs:
                hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = color

        document.add_heading("Findings", level=1)
        
        for i, finding in enumerate(result.findings):
            document.add_heading(f"Finding {i+1}: {finding.description}", level=2)
            
            p = document.add_paragraph()
            p.add_run("Status: ").bold = True
            status_run = p.add_run(finding.status)
            if finding.status == "Full":
                status_run.font.color.rgb = RGBColor(0, 128, 0)
            elif finding.status == "Partial":
                status_run.font.color.rgb = RGBColor(218, 165, 32)
            else:
                status_run.font.color.rgb = RGBColor(255, 0, 0)
            
            document.add_paragraph(f"Reasoning: {finding.reasoning}")
            
            if finding.baseline_match_text:
                document.add_paragraph("Baseline Match:", style='Intense Quote')
                document.add_paragraph(finding.baseline_match_text)

            if finding.draft_amendment:
                amendment_header = document.add_paragraph()
                amendment_header.add_run("Suggested Amendment: ").bold = True
                amendment_paragraph = document.add_paragraph(finding.draft_amendment)
                amendment_paragraph.paragraph_format.left_indent = Inches(0.25)

        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)
        return buf.read()

    try:
        docx_bytes = await asyncio.to_thread(create_docx)

        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="lacuna-gap-report.docx"'
            },
        )
    except Exception as e:
        logger.error(f"Error generating Word report: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to generate Word report")


@router.post("/gap-analysis/batch", response_model=BatchGapAnalysisResponse)
async def batch_gap_analysis(
    request: BatchGapAnalysisRequest, service=Depends(get_gap_analysis_service)
):
    """
    Perform gap analysis for one baseline against multiple circulars in parallel.
    Individual failures are returned per circular without failing the whole batch.
    """

    async def _analyze_single(circular_doc_id: str) -> BatchGapAnalysisResult:
        cache_key = (
            circular_doc_id,
            request.baseline_id,
            request.is_policy_baseline,
            request.include_amendments,
        )
        if cache_key in _gap_cache and not request.no_llm:
            logger.info(
                f"Gap analysis cache hit — returning cached result for {circular_doc_id}"
            )
            return BatchGapAnalysisResult(
                circular_doc_id=circular_doc_id, result=_gap_cache[cache_key]
            )

        try:
            result = await service.perform_gap_analysis(
                circular_doc_id=circular_doc_id,
                baseline_id=request.baseline_id,
                is_policy_baseline=request.is_policy_baseline,
                include_amendments=request.include_amendments,
                no_llm=request.no_llm,
            )
            if not request.no_llm:
                _gap_cache[cache_key] = result
            return BatchGapAnalysisResult(circular_doc_id=circular_doc_id, result=result)
        except ValueError as e:
            return BatchGapAnalysisResult(circular_doc_id=circular_doc_id, error=str(e))
        except Exception as e:
            logger.error(
                f"Error performing gap analysis for {circular_doc_id}: {e}",
                exc_info=True,
            )
            return BatchGapAnalysisResult(
                circular_doc_id=circular_doc_id, error="Internal server error"
            )

    tasks = [_analyze_single(circular_doc_id) for circular_doc_id in request.circular_doc_ids]
    gathered_results = await asyncio.gather(*tasks, return_exceptions=True)

    results = []
    for circular_doc_id, item in zip(request.circular_doc_ids, gathered_results):
        if isinstance(item, Exception):
            logger.error(
                f"Error performing gap analysis for {circular_doc_id}: {item}",
                exc_info=True,
            )
            results.append(
                BatchGapAnalysisResult(
                    circular_doc_id=circular_doc_id, error="Internal server error"
                )
            )
        else:
            results.append(item)

    return BatchGapAnalysisResponse(baseline_id=request.baseline_id, results=results)
